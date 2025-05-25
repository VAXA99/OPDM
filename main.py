import streamlit as st
import docx
from docx import Document
import openpyxl
from openpyxl import load_workbook
from openpyxl.utils.datetime import from_excel
from pathlib import Path
import tempfile
import shutil
import os
import logging
import zipfile
import datetime
import re

logging.basicConfig(level=logging.INFO)

def detect_columns(sheet):
    cols = {"fio": None, "dob": None, "position": None, "risk": None, "diagnosis": None, "header": None}
    for r in range(1, 21):
        for c in range(1, 40):
            val = sheet.cell(row=r, column=c).value
            if not isinstance(val, str):
                continue
            txt = val.strip().lower()
            if "фио" in txt and "сотрудник" in txt:
                cols["fio"] = c
                cols["header"] = r
            if "дата" in txt and "рожд" in txt:
                cols["dob"] = c
                cols["header"] = cols["header"] or r
            if "штатная должность" in txt:
                cols["position"] = c
            if "факторы риска" in txt:
                cols["risk"] = c
            if "мкб-10" in txt:
                cols["diagnosis"] = c
        if all(cols.values()):
            break
    return cols if all(cols.values()) else None

def excel_date_to_str(value):
    if value is None:
        return ""
    if isinstance(value, datetime.datetime):
        return value.strftime("%d.%m.%Y")
    if isinstance(value, (int, float)):
        try:
            return from_excel(value).strftime("%d.%m.%Y")
        except:
            return str(value)
    return str(value)

def make_safe_filename(text):
    return re.sub(r"[^\w\-_. ]", "_", text.strip())

def get_downloads_folder():
    if os.name == "nt":
        try:
            import winreg
            sub_key = r'SOFTWARE\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders'
            downloads_guid = '{374DE290-123F-4565-9164-39C4925E467B}'
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, sub_key) as key:
                downloads_path = winreg.QueryValueEx(key, downloads_guid)[0]
            return Path(downloads_path)
        except Exception:
            return Path.home() / "Downloads"
    else:
        return Path.home() / "Downloads"

def get_desktop_folder():
    if os.name == "nt":
        try:
            import winreg
            sub_key = r'SOFTWARE\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders'
            desktop_guid = 'Desktop'
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, sub_key) as key:
                desktop_path = winreg.QueryValueEx(key, desktop_guid)[0]
            return Path(desktop_path)
        except Exception:
            return Path.home() / "Desktop"
    else:
        return Path.home() / "Desktop"

def get_all_drives():
    drives = []
    if os.name == 'nt':
        import string
        from ctypes import windll
        bitmask = windll.kernel32.GetLogicalDrives()
        for letter in string.ascii_uppercase:
            if bitmask & 1:
                drives.append(f"{letter}:\\")
            bitmask >>= 1
    else:
        drives.append("/")
        media = Path("/media")
        if media.exists():
            drives.extend([str(p) for p in media.iterdir() if p.is_dir()])
        mnt = Path("/mnt")
        if mnt.exists():
            drives.extend([str(p) for p in mnt.iterdir() if p.is_dir()])
    return drives

def build_fs_locations():
    locations = []
    locations.append(("Загрузки", get_downloads_folder()))
    locations.append(("Рабочий стол", get_desktop_folder()))
    for d in get_all_drives():
        label = f"Диск {d[0]}"
        locations.append((label, Path(d)))
    return locations

st.set_page_config(page_title="Генерация Word-документов", layout="centered")
st.title("🔧 Генерация персональных Word-документов")

custom_jobplace = st.text_input("💼 Введите место работы:", value="ГБОУ Школа №")

save_to_fs = st.checkbox("💾 Сохранять DOCX-файлы в файловой системе", value=True)
save_zip_to_fs = st.checkbox("💾 Сохранять ZIP-архив в файловой системе", value=True)

# Для DOCX
if save_to_fs:
    locations = build_fs_locations()
    location_labels = [f"{name} — {str(path)}" for name, path in locations]
    default_idx = 0
    selected_idx = st.selectbox("🌍 Куда сохранить DOCX-файлы:",
                                options=list(range(len(location_labels))),
                                format_func=lambda i: location_labels[i],
                                index=default_idx)
    selected_path = locations[selected_idx][1]
    docx_subdir = st.text_input("Подпапка для DOCX-файлов (создастся автоматически, оставьте пустым для корня):", value="generated_docs", key="docx_subdir")
    if docx_subdir:
        target_dir = selected_path / docx_subdir
    else:
        target_dir = selected_path
    st.info(f"DOCX-файлы сохранятся в: {target_dir}")

# Для ZIP
if save_zip_to_fs:
    zip_locations = build_fs_locations()
    zip_location_labels = [f"{name} — {str(path)}" for name, path in zip_locations]
    zip_default_idx = 0
    zip_selected_idx = st.selectbox("🌍 Куда сохранить ZIP-архив:",
                                    options=list(range(len(zip_location_labels))),
                                    format_func=lambda i: zip_location_labels[i],
                                    index=zip_default_idx,
                                    key="zip_fs_selectbox")
    zip_selected_path = zip_locations[zip_selected_idx][1]
    zip_filename = st.text_input("Имя ZIP-архива (без .zip):", value="generated_docs", key="zip_filename")
    zip_final_save_dir = st.text_input("Подпапка для ZIP-архива (оставьте пустым для корня):", value="", key="zip_subdir")
    if zip_final_save_dir:
        zip_target_dir = zip_selected_path / zip_final_save_dir
    else:
        zip_target_dir = zip_selected_path
    zip_full_path = zip_target_dir / f"{zip_filename}.zip"
    st.info(f"ZIP-архив сохранится в: {zip_full_path}")

excel_file = st.file_uploader("📄 Загрузите Excel-файл с данными", type=["xlsx"])
word_template = st.file_uploader("📅 Выберите шаблон Word", type=["docx"])

save_fs_success = None
zip_fs_success = None

if excel_file and word_template and st.button("✅ Начать генерацию"):
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            excel_path = Path(tmpdir) / "data.xlsx"
            template_path = Path(tmpdir) / "template.docx"
            output_path = Path(tmpdir) / "docs"
            output_path.mkdir(exist_ok=True)

            with open(excel_path, "wb") as f:
                f.write(excel_file.read())
            with open(template_path, "wb") as f:
                f.write(word_template.read())

            wb = load_workbook(excel_path, data_only=True)
            sheet = wb.active

            cols = detect_columns(sheet)
            if not cols:
                st.error("❌ Неверный шаблон Excel: не найдены нужные колонки.")
            else:
                header_row = cols["header"] + 1

                with st.spinner("⏳ Подождите, идёт генерация..."):
                    total_rows = sheet.max_row - header_row + 1
                    progress_bar = st.progress(0)
                    counter = 0

                    for row in range(header_row, sheet.max_row + 1):
                        fio = str(sheet.cell(row=row, column=cols["fio"]).value or "").strip()
                        if not fio:
                            continue
                        dob_val = sheet.cell(row=row, column=cols["dob"]).value
                        birthdate = excel_date_to_str(dob_val)
                        position = str(sheet.cell(row=row, column=cols["position"]).value or "").strip()
                        risk = str(sheet.cell(row=row, column=cols["risk"]).value or "").strip()
                        diagnosis = str(sheet.cell(row=row, column=cols["diagnosis"]).value or "").strip()

                        dest_file = output_path / (make_safe_filename(fio) + ".docx")
                        shutil.copy(template_path, dest_file)
                        doc = Document(dest_file)

                        replacements = {
                            "1. Ф.И.О": f"1. Ф.И.О: {fio} {birthdate} г.р.",
                            "2. Место работы": f"2. Место работы: {custom_jobplace}",
                            "3. Профессия (должность) (в настоящее время)": f"3. Профессия (должность) (в настоящее время): {position}",
                            "Вредный производственный фактор": f"Вредный производственный фактор, наименование вида работ: {risk}",
                            "6. Наименование": f"6. Наименование: {diagnosis}"
                        }

                        for p in doc.paragraphs:
                            for k, v in replacements.items():
                                if p.text.strip().startswith(k):
                                    for run in p.runs:
                                        run.text = ""
                                    if p.runs:
                                        p.runs[0].text = v

                        for table in doc.tables:
                            for row_cell in table.rows:
                                for cell in row_cell.cells:
                                    for p in cell.paragraphs:
                                        for k, v in replacements.items():
                                            if p.text.strip().startswith(k):
                                                for run in p.runs:
                                                    run.text = ""
                                                if p.runs:
                                                    p.runs[0].text = v

                        doc.save(dest_file)
                        counter += 1
                        progress_bar.progress(min(counter / total_rows, 1.0))

                st.success(f"✅ Документы успешно созданы: {counter} файл(ов)")

                # --- Сохраняем DOCX ---
                if save_to_fs:
                    with st.spinner("⏳ Копируем DOCX-файлы в файловую систему..."):
                        total_docs = len(list(output_path.glob("*.docx")))
                        fs_bar = st.progress(0)
                        done = 0
                        target_dir.mkdir(parents=True, exist_ok=True)
                        for docx_file in output_path.glob("*.docx"):
                            shutil.copy(docx_file, target_dir / docx_file.name)
                            done += 1
                            fs_bar.progress(min(done / total_docs, 1.0))
                    save_fs_success = str(target_dir)

                # --- Сохраняем ZIP ---
                if save_zip_to_fs:
                    from io import BytesIO
                    zip_path = Path(tmpdir) / f"{zip_filename}.zip"
                    with zipfile.ZipFile(zip_path, "w") as zipf:
                        for docx_file in output_path.glob("*.docx"):
                            zipf.write(docx_file, arcname=docx_file.name)
                    with open(zip_path, "rb") as f:
                        zip_buffer = f.read()
                    zip_target_dir.mkdir(parents=True, exist_ok=True)
                    zip_full_path = zip_target_dir / f"{zip_filename}.zip"
                    with open(zip_full_path, "wb") as f:
                        f.write(zip_buffer)
                    zip_fs_success = str(zip_full_path)

        except Exception as e:
            st.exception(e)

if save_to_fs and save_fs_success:
    st.success(f"DOCX-файлы сохранены в: {save_fs_success}")

if save_zip_to_fs and zip_fs_success:
    st.success(f"ZIP-архив сохранён в: {zip_fs_success}")
