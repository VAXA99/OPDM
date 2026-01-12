import streamlit as st
from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.datetime import from_excel
from pathlib import Path
import tempfile
import shutil
import os
import logging
import datetime
import re

logging.basicConfig(level=logging.INFO)


# ------------------------
# Excel parsing
# ------------------------
def detect_columns(sheet):
    cols = {
        "fio": None,
        "dob": None,
        "position": None,
        "risk": None,
        "diagnosis": None,
        "address": None,
        "header": None,
    }

    for r in range(1, 21):
        for c in range(1, 80):
            val = sheet.cell(row=r, column=c).value
            if not isinstance(val, str):
                continue

            txt = val.strip().lower()

            if "фио" in txt:
                cols["fio"] = cols["fio"] or c
                cols["header"] = cols["header"] or r

            if ("дата" in txt and "рожд" in txt) or ("д.р" in txt) or ("д р" in txt) or ("др" == txt.replace(".", "").replace(" ", "")):
                cols["dob"] = cols["dob"] or c
                cols["header"] = cols["header"] or r

            if "адрес" in txt:
                cols["address"] = cols["address"] or c
                cols["header"] = cols["header"] or r

            if "штатная должность" in txt or ("должность" in txt and "штат" in txt):
                cols["position"] = cols["position"] or c
            if "факторы риска" in txt or ("фактор" in txt and "риска" in txt):
                cols["risk"] = cols["risk"] or c
            if "мкб-10" in txt or "мкб 10" in txt or "мкб10" in txt:
                cols["diagnosis"] = cols["diagnosis"] or c

    return cols


def validate_columns(cols: dict, mode: str):
    if mode == "Заключение предварительное":
        required = ["fio", "dob", "position", "risk", "diagnosis", "header"]
    else:
        required = ["fio", "dob", "address", "header"]
    return [k for k in required if not cols.get(k)]


def excel_date_to_str(value):
    if value is None:
        return ""
    if isinstance(value, datetime.datetime):
        return value.strftime("%d.%m.%Y")
    if isinstance(value, datetime.date):
        return datetime.datetime(value.year, value.month, value.day).strftime("%d.%m.%Y")
    if isinstance(value, (int, float)):
        try:
            return from_excel(value).strftime("%d.%m.%Y")
        except Exception:
            return str(value)
    return str(value)


def make_safe_filename(text: str) -> str:
    return re.sub(r"[^\w\-_. ]", "_", text.strip())


# ------------------------
# FS locations
# ------------------------
def get_downloads_folder():
    if os.name == "nt":
        try:
            import winreg
            sub_key = r"SOFTWARE\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders"
            downloads_guid = "{374DE290-123F-4565-9164-39C4925E467B}"
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, sub_key) as key:
                downloads_path = winreg.QueryValueEx(key, downloads_guid)[0]
            return Path(downloads_path)
        except Exception:
            return Path.home() / "Downloads"
    return Path.home() / "Downloads"


def get_desktop_folder():
    if os.name == "nt":
        try:
            import winreg
            sub_key = r"SOFTWARE\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders"
            desktop_guid = "Desktop"
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, sub_key) as key:
                desktop_path = winreg.QueryValueEx(key, desktop_guid)[0]
            return Path(desktop_path)
        except Exception:
            return Path.home() / "Desktop"
    return Path.home() / "Desktop"


def get_all_drives():
    drives = []
    if os.name == "nt":
        import string
        from ctypes import windll
        bitmask = windll.kernel32.GetLogicalDrives()
        for letter in string.ascii_uppercase:
            if bitmask & 1:
                drives.append(f"{letter}:\\")
            bitmask >>= 1
    else:
        drives.append("/")
    return drives


def build_fs_locations():
    locations = [
        ("Загрузки", get_downloads_folder()),
        ("Рабочий стол", get_desktop_folder()),
    ]
    for d in get_all_drives():
        label = f"Диск {d[0]}"
        locations.append((label, Path(d)))
    return locations


# ------------------------
# DOCX replacement helpers
# ------------------------
def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _is_field_run(run) -> bool:
    """
    Определяем run, который является "полем" с линией.
    Для твоего шаблона это обычно:
    - underline=True и текст пустой/пробелы
    - или очень много пробелов/подчёркиваний
    """
    t = run.text or ""
    stripped = t.strip()

    # если подчёркнут и это пробелы/пустота — это почти точно поле
    if getattr(run.font, "underline", None):
        if stripped == "" and len(t) >= 2:
            return True

    # поле может быть без underline на самом run (underline в стиле),
    # но текст тогда длинный и пробельный
    if stripped == "" and len(t) >= 8:
        return True

    # поле из "_" (редко, но поддержим)
    if t.count("_") >= 5:
        return True

    return False


def fill_underlined_field_run(field_run, value: str):
    """
    Заполняем подчёркнутое поле:
    value + хвост NBSP до исходной длины run.
    NBSP нужен, чтобы Word не схлопывал пробелы.
    """
    original = field_run.text or ""
    original_len = len(original)

    NBSP = "\u00A0"

    value = value or ""

    # Гарантируем подчёркивание на поле (если оно было на стиле/частично, это не мешает)
    field_run.font.underline = True

    # Если value длиннее поля — не режем молча: поле расширится.
    # Если хочешь строго по длине, раскомментируй:
    # if len(value) > original_len:
    #     value = value[:original_len]

    pad_len = max(0, original_len - len(value))
    field_run.text = value + (NBSP * pad_len)


def replace_in_paragraph_keep_format(paragraph, key_predicate, new_value, once_state_key=None, once_state=None):
    """
    Для "диаскин":
    - находим строку по key_predicate
    - находим run-поле (подчёркнутые пробелы)
    - заполняем его value + NBSP-хвостом, подчёркивание остаётся
    """
    full_text = paragraph.text or ""
    if not key_predicate(full_text):
        return False

    low = full_text.lower()
    if ("подпись" in low) or ("печать" in low) or ("направившего" in low):
        return False

    if once_state_key and once_state is not None and once_state.get(once_state_key, False):
        return False

    field_run = None
    for run in paragraph.runs:
        if _is_field_run(run):
            field_run = run
            break

    if field_run is None:
        # fallback: заменить после двоеточия (может убрать линию, но лучше чем ничего)
        if ":" in full_text:
            prefix = full_text.split(":", 1)[0] + ":"
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = f"{prefix} {new_value}"
            else:
                paragraph.add_run(f"{prefix} {new_value}")
            if once_state_key and once_state is not None:
                once_state[once_state_key] = True
            return True
        return False

    fill_underlined_field_run(field_run, str(new_value or ""))

    if once_state_key and once_state is not None:
        once_state[once_state_key] = True
    return True


def apply_replacements_old_logic(doc: Document, custom_jobplace: str, fio: str, birthdate: str, position: str, risk: str, diagnosis: str):
    replacements = {
        "1. Ф.И.О": f"1. Ф.И.О: {fio} {birthdate} г.р.",
        "2. Место работы": f"2. Место работы: {custom_jobplace}",
        "3. Профессия (должность) (в настоящее время)": f"3. Профессия (должность) (в настоящее время): {position}",
        "Вредный производственный фактор": f"Вредный производственный фактор, наименование вида работ: {risk}",
        "6. Наименование": f"6. Наименование: {diagnosis}",
    }

    for p in iter_all_paragraphs(doc):
        txt = (p.text or "").strip()
        for k, v in replacements.items():
            if txt.startswith(k):
                for run in p.runs:
                    run.text = ""
                if p.runs:
                    p.runs[0].text = v
                else:
                    p.add_run(v)
                break


def apply_replacements_diaskin(doc: Document, fio: str, birthdate: str, address: str):
    once = {"fio_done": False, "dob_done": False, "addr_done": False}

    for p in iter_all_paragraphs(doc):

        def is_fio_top(t: str) -> bool:
            t2 = (t or "").replace(" ", "")
            return ("ф.и.о.:" in t2.lower()) or ("фио:" in t2.lower())

        replace_in_paragraph_keep_format(
            p,
            key_predicate=is_fio_top,
            new_value=fio,
            once_state_key="fio_done",
            once_state=once,
        )

        def is_dob_line(t: str) -> bool:
            tl = (t or "").lower()
            return ("дата" in tl) and ("рожд" in tl)

        if not once["dob_done"]:
            replace_in_paragraph_keep_format(
                p,
                key_predicate=is_dob_line,
                new_value=birthdate,
                once_state_key="dob_done",
                once_state=once,
            )

        def is_addr_line(t: str) -> bool:
            tl = (t or "").lower()
            return ("адрес" in tl) and ("постоянного" in tl) and ("житель" in tl)

        if not once["addr_done"]:
            replace_in_paragraph_keep_format(
                p,
                key_predicate=is_addr_line,
                new_value=address,
                once_state_key="addr_done",
                once_state=once,
            )


# ------------------------
# UI
# ------------------------
st.set_page_config(page_title="ПРОФПАК", layout="centered")
st.title("ПРОФПАК")

mode = st.radio("📄 Тип документа", ["Заключение предварительное", "Направление на диаскин"], horizontal=True)

custom_jobplace = st.text_input("💼 Введите место работы:", value="ГБОУ Школа №")

save_to_fs = st.checkbox("💾 Сохранять DOCX-файлы в файловой системе", value=True)

target_dir = None
if save_to_fs:
    locations = build_fs_locations()
    location_labels = [f"{name} — {str(path)}" for name, path in locations]
    selected_idx = st.selectbox(
        "🌍 Куда сохранить DOCX-файлы:",
        options=list(range(len(location_labels))),
        format_func=lambda i: location_labels[i],
        index=0,
    )
    selected_path = locations[selected_idx][1]
    docx_subdir = st.text_input(
        "Подпапка для DOCX-файлов (создастся автоматически, оставьте пустым для корня):",
        value="generated_docs",
        key="docx_subdir",
    )
    target_dir = selected_path / docx_subdir if docx_subdir else selected_path
    st.info(f"DOCX-файлы сохранятся в: {target_dir}")

excel_file = st.file_uploader("📄 Загрузите Excel-файл с данными", type=["xlsx"])
word_template = st.file_uploader("📄 Выберите шаблон Word", type=["docx"])

save_fs_success = None

if excel_file and word_template and st.button("✅ Начать генерацию"):
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            tmpdir = Path(tmpdir)

            excel_path = tmpdir / "data.xlsx"
            template_path = tmpdir / "template.docx"
            output_path = tmpdir / "docs"
            output_path.mkdir(exist_ok=True)

            excel_path.write_bytes(excel_file.read())
            template_path.write_bytes(word_template.read())

            wb = load_workbook(excel_path, data_only=True)
            sheet = wb.active

            cols = detect_columns(sheet)
            missing = validate_columns(cols, mode)
            if missing:
                st.error("❌ Неверный Excel-шаблон для выбранного режима.\n\n" f"Не найдены колонки: {', '.join(missing)}")
                st.stop()

            header_row = (cols["header"] or 1) + 1

            with st.spinner("⏳ Подождите, идёт генерация..."):
                total_rows = max(sheet.max_row - header_row + 1, 1)
                progress_bar = st.progress(0)

                counter = 0
                processed = 0

                for r in range(header_row, sheet.max_row + 1):
                    processed += 1

                    fio = str(sheet.cell(row=r, column=cols["fio"]).value or "").strip()
                    if not fio:
                        progress_bar.progress(min(processed / total_rows, 1.0))
                        continue

                    dob_val = sheet.cell(row=r, column=cols["dob"]).value
                    birthdate = excel_date_to_str(dob_val)

                    address = ""
                    if cols.get("address"):
                        address = str(sheet.cell(row=r, column=cols["address"]).value or "").strip()

                    position = risk = diagnosis = ""
                    if mode == "Заключение предварительное":
                        position = str(sheet.cell(row=r, column=cols["position"]).value or "").strip()
                        risk = str(sheet.cell(row=r, column=cols["risk"]).value or "").strip()
                        diagnosis = str(sheet.cell(row=r, column=cols["diagnosis"]).value or "").strip()

                    dest_file = output_path / (make_safe_filename(fio) + ".docx")
                    shutil.copy(template_path, dest_file)
                    doc = Document(dest_file)

                    if mode == "Заключение предварительное":
                        apply_replacements_old_logic(doc, custom_jobplace, fio, birthdate, position, risk, diagnosis)
                    else:
                        apply_replacements_diaskin(doc, fio, birthdate, address)

                    doc.save(dest_file)
                    counter += 1
                    progress_bar.progress(min(processed / total_rows, 1.0))

            st.success(f"✅ Документы успешно созданы: {counter} файл(ов)")

            if save_to_fs and target_dir is not None:
                docs = list(output_path.glob("*.docx"))
                if not docs:
                    st.warning("⚠️ DOCX-файлы не были созданы (возможно, пустые строки/нет ФИО).")
                else:
                    with st.spinner("⏳ Копируем DOCX-файлы в файловую систему..."):
                        target_dir.mkdir(parents=True, exist_ok=True)
                        fs_bar = st.progress(0)
                        for i, docx_file in enumerate(docs, start=1):
                            shutil.copy(docx_file, target_dir / docx_file.name)
                            fs_bar.progress(i / len(docs))
                    save_fs_success = str(target_dir)

        except Exception as e:
            st.exception(e)

if save_to_fs and save_fs_success:
    st.success(f"DOCX-файлы сохранены в: {save_fs_success}")
